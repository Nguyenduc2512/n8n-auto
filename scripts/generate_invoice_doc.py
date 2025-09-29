import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from datetime import datetime
import sys, json, base64, unicodedata, re

# Thư mục chứa ảnh (dùng đúng path trong container)
folder = Path("/data/outputs")

def to_camel_ascii(s: str) -> str:
    if not s:
        return ''

    # 1. Thay đ → d, Đ → D trước
    s = s.replace('đ', 'd').replace('Đ', 'D')

    # 2. Chuẩn hóa unicode để tách dấu (loại dấu tiếng Việt)
    s = unicodedata.normalize('NFD', s)
    s = ''.join(c for c in s if not unicodedata.combining(c))

    # 3. Loại ký tự không phải chữ/số, chuyển về khoảng trắng
    s = re.sub(r'[^a-zA-Z0-9]+', ' ', s)

    # 4. Cắt và camelCase
    parts = s.strip().split()
    if not parts:
        return ''

    camel = parts[0].lower() + ''.join(p.capitalize() for p in parts[1:])

    # 5. Thêm "_" nếu bắt đầu bằng số
    if camel[0].isdigit():
        camel = '_' + camel

    return camel


def read_data():
    """Đọc JSON từ stdin (n8n gửi vào)."""
    if '--b64' in sys.argv:
        b64 = sys.argv[sys.argv.index('--b64') + 1]
        data = json.loads(base64.b64decode(b64).decode('utf-8'))
    else:
        data = json.load(sys.stdin)

    return data

data = read_data()


new_name = None
for item in data:
    if "newName" in item:
        new_name = item["newName"]
        break
if not new_name:
    new_name = "screenshots"

safe_name = to_camel_ascii(new_name) or "screenshots"
output_file = folder / f"{safe_name}.docx"

# Lấy danh sách file ảnh, bỏ qua captcha
files = [
    f for f in folder.iterdir()
    if f.suffix.lower() in [".png", ".jpg", ".jpeg"] and "_captcha" not in f.name
]

# Sort theo thời gian chỉnh sửa (gần nhất với thời gian tạo trong container)
files.sort(key=lambda f: f.stat().st_mtime)

# Tạo docx
doc = Document()

for f in files:
    created = datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S")
    doc.add_picture(str(f), width=Inches(5))  # scale ảnh vừa trang
    doc.add_paragraph("")  # dòng trống

doc.save(output_file)
print(f"✅ Saved {output_file}")
